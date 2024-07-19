import os
import akshare as ak
import pandas as pd
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt
from matplotlib import rcParams
import mplfinance as mpf
import streamlit as st

# 使用系统字体 SimHei
rcParams['font.sans-serif'] = ['SimHei']
rcParams['axes.unicode_minus'] = False

# 创建文件夹和文档保存路径
def create_folder_and_doc_path(custom_date):
    base_path = "C:/Users/jacky/Desktop/期货日报"
    folder_path = os.path.join(base_path, f"期货日报_{custom_date}")
    if not os.path.exists(folder_path):
        os.makedirs(folder_path)
    base_filename = "恒力期货日报"
    filename = f"{base_filename}_{custom_date}.docx"
    doc_path = os.path.join(folder_path, filename)
    return doc_path, folder_path

# 设置文档样式
def set_doc_style(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "楷体")
    normal.font.size = Pt(12)

# 获取当天行情概述数据
def get_market_trend_data(symbol, custom_date):
    try:
        today = custom_date
        yesterday = today - timedelta(days=1)
        start_time = yesterday.strftime('%Y-%m-%d') + ' 21:00:00'
        end_time = today.strftime('%Y-%m-%d') + ' 23:00:00'
        df = ak.futures_zh_minute_sina(symbol=symbol, period="1")
        df['datetime'] = pd.to_datetime(df['datetime'])
        filtered_data = df[(df['datetime'] >= start_time) & (df['datetime'] <= end_time)]
        
        if filtered_data.empty:
            return "", "", pd.DataFrame()

        # 获取开盘价和收盘价
        day_open_price = filtered_data.iloc[0]['open']  # 前一晚21:00开盘价
        day_close_price = filtered_data[filtered_data['datetime'] <= today.strftime('%Y-%m-%d') + ' 15:00:00'].iloc[-1]['close']  # 15:00收盘价

        high_price = filtered_data['high'].max()
        low_price = filtered_data['low'].min()
        price_change = day_close_price - day_open_price
        price_change_percentage = (price_change / day_open_price) * 100
        trend = "上涨" if price_change > 0 else "下跌" if price_change < 0 else "持平"
        day_description = (
            f"{custom_date.strftime('%Y-%m-%d')}日{symbol}主力合约开盘价为{day_open_price}元/吨，最高价为{high_price}元/吨，"
            f"最低价为{low_price}元/吨，收盘价为{day_close_price}元/吨，较前一日{trend}了"
            f"{abs(price_change):.2f}元/吨，涨跌幅为{price_change_percentage:.2f}%。"
        )

        # 获取夜盘走势
        night_start_time = today.strftime('%Y-%m-%d') + ' 21:00:00'
        night_end_time = (today + timedelta(days=1)).strftime('%Y-%m-%d') + ' 01:00:00'
        night_data = df[(df['datetime'] >= night_start_time) & (df['datetime'] <= night_end_time)]
        if night_data.empty:
            night_description = "夜盘数据不可用。"
        else:
            night_open_price = night_data.iloc[0]['open']
            night_close_price = night_data.iloc[-1]['close']
            night_price_change = night_close_price - night_open_price
            night_price_change_percentage = (night_price_change / night_open_price) * 100
            night_trend = "上涨" if night_price_change > 0 else "下跌" if night_price_change < 0 else "持平"
            night_description = (
                f"夜盘走势：开盘价为{night_open_price}元/吨，收盘价为{night_close_price}元/吨，较开盘{night_trend}了"
                f"{abs(night_price_change):.2f}元/吨，涨跌幅为{night_price_change_percentage:.2f}%。"
            )

        return day_description, night_description, filtered_data
    except Exception as e:
        return f"获取市场走势数据失败: {e}", "", pd.DataFrame()

# 创建K线图
def create_k_line_chart(data, symbol, folder_path):
    if data.empty:
        print("数据为空，无法生成K线图。")
        return None
    data.set_index('datetime', inplace=True)
    data = data[['open', 'high', 'low', 'close']]
    data.columns = ['Open', 'High', 'Low', 'Close']
    fig, ax = plt.subplots(figsize=(10, 6))
    mpf.plot(data, type='candle', style='charles', ax=ax)
    k_line_chart_path = os.path.join(folder_path, 'k_line_chart.png')
    plt.savefig(k_line_chart_path)
    plt.close(fig)
    return k_line_chart_path

# 获取新闻数据
def get_news_data(symbol):
    try:
        symbol_mapping = {
            'cu': '铜',
            'al': '铝',
            'pb': '铅',
            'zn': '锌',
            'ni': '镍',
            'sn': '锡'
        }
        symbol_name = symbol_mapping.get(symbol[:2], '未知品种')
        
        df = ak.futures_news_shmet(symbol=symbol_name)
        df['发布时间'] = pd.to_datetime(df['发布时间']).dt.tz_localize('Asia/Shanghai')
        
        latest_news = df.tail(20)
        description = ""
        for index, row in latest_news.iterrows():
            description += f"{row['发布时间'].strftime('%Y-%m-%d %H:%M:%S %Z')} - {row['内容']}\n"
        return description
    except Exception as e:
        return f"获取新闻数据失败: {e}"

# 设置楷体字体
def set_font_kaiti(paragraph):
    run = paragraph.add_run()
    run.font.name = '楷体'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
    run.font.size = Pt(12)
    return run

# 创建报告
def create_report(custom_date_str, symbol, user_description, main_view):
    custom_date = datetime.strptime(custom_date_str, '%Y-%m-%d')
    doc_path, folder_path = create_folder_and_doc_path(custom_date_str)
    market_trend_description, night_trend_description, market_data = get_market_trend_data(symbol=symbol, custom_date=custom_date)
    
    if market_data.empty:
        st.error("无法生成报告，因为市场数据为空。")
        return None
    
    news_description = get_news_data(symbol)
    
    k_line_chart_path = create_k_line_chart(market_data, symbol, folder_path)

    doc = Document()
    set_doc_style(doc)

    # 添加标题
    title = doc.add_paragraph(f"恒力期货日报 {custom_date_str}")
    title_run = title.runs[0]
    title_run.font.size = Pt(14)
    title_run.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加品种名
    commodity_name = {
        'cu': '铜',
        'al': '铝',
        'pb': '铅',
        'zn': '锌',
        'ni': '镍',
        'sn': '锡'
    }.get(symbol[:2], '未知品种')
    
    commodity_paragraph = doc.add_paragraph(commodity_name)
    commodity_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加主要观点段落
    main_view_paragraph = doc.add_paragraph()
    main_view_run = main_view_paragraph.add_run("主要观点：")
    main_view_run.bold = True
    main_view_paragraph.add_run("\n" + main_view)

    # 添加核心逻辑段落
    core_logic = doc.add_paragraph()
    core_logic_run = core_logic.add_run("核心逻辑：")
    core_logic_run.bold = True
    core_logic.add_run("\n")

    # 添加昨日走势
    market_trend_paragraph = doc.add_paragraph()
    market_trend_run = market_trend_paragraph.add_run("昨日走势：")
    market_trend_run.bold = True

    if k_line_chart_path:
        doc.add_picture(k_line_chart_path, width=Inches(6))
    market_trend_paragraph.add_run("\n" + user_description)
    set_font_kaiti(market_trend_paragraph)

    # 添加今日新闻资讯
    news_paragraph = doc.add_paragraph()
    news_run = news_paragraph.add_run("今日新闻资讯：")
    news_run.bold = True
    set_font_kaiti(news_paragraph)
    news_paragraph.add_run("\n" + news_description)

    # 添加附录
    appendix = doc.add_paragraph()
    appendix_run = appendix.add_run("附录")
    appendix_run.bold = True

    # 保存文档
    doc.save(doc_path)
    return doc_path

# Streamlit应用
st.title("期货日报生成器")
st.write("created by 恒力期货上海分公司")

custom_date = st.date_input("请选择日期")
symbol = st.selectbox("请选择品种", ['cu', 'al', 'pb', 'zn', 'ni', 'sn'])
full_contract = st.text_input("请输入完整品种合约（如：CU2408）")

if st.button("生成K线图"):
    custom_date_str = custom_date.strftime('%Y-%m-%d')
    day_description, night_description, market_data = get_market_trend_data(full_contract, custom_date)
    k_line_chart_path = create_k_line_chart(market_data, full_contract, ".")

    if k_line_chart_path:
        st.image(k_line_chart_path, caption="昨日K线图")
    else:
        st.error("无法生成K线图，因为市场数据为空。")
    
    st.write("昨日走势：")
    st.write(day_description)
    st.write(night_description)

user_description = st.text_area("请输入行情描述（自动生成或自行编辑）")
main_view = st.text_area("请输入主要观点")

if st.button("生成日报"):
    custom_date_str = custom_date.strftime('%Y-%m-%d')
    doc_path = create_report(custom_date_str, full_contract, user_description, main_view)
    if doc_path:
        with open(doc_path, "rb") as f:
            st.download_button(
                label="下载日报",
                data=f,
                file_name=os.path.basename(doc_path),
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
